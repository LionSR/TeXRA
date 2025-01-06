import argparse

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE


def apply_styles(document):
    styles = document.styles

    font_name = "Arial"

    # Define styles
    styles.add_style("Cover Letter", WD_STYLE_TYPE.PARAGRAPH)
    styles["Cover Letter"].font.name = font_name
    styles["Cover Letter"].font.size = Pt(11)

    styles.add_style("Response Heading", WD_STYLE_TYPE.PARAGRAPH)
    styles["Response Heading"].font.name = font_name
    styles["Response Heading"].font.size = Pt(11)
    styles["Response Heading"].font.bold = True

    styles.add_style("Comment", WD_STYLE_TYPE.PARAGRAPH)
    styles["Comment"].font.name = font_name
    styles["Comment"].font.size = Pt(11)
    styles["Comment"].font.italic = True

    styles.add_style("Response", WD_STYLE_TYPE.PARAGRAPH)
    styles["Response"].font.name = font_name
    styles["Response"].font.size = Pt(11)
    styles["Response"].font.color.rgb = RGBColor(0, 0, 255)  # Set response text to blue


def process_text(text, document):
    lines = text.replace("\n\n", "\n").split("\n")
    cover_letter_has_ended = False
    referee_comment_has_started = False
    for line in lines:
        if "<reply_letter>" in line:
            pass
        elif "<cover_letter>" in line:
            document.add_paragraph("Cover Letter", style="Response Heading")
        elif line.startswith("</cover_letter>"):
            cover_letter_has_ended = True
            document.add_page_break()
            pass
        elif line.startswith("<point_by_point_reply_"):
            document.add_paragraph(
                "-----------------------------------------------------------------------",
                style="Response Heading",
            )
            document.add_paragraph(f"Responses to Referee {line[-2].upper()}", style="Response Heading")
        elif line.startswith("</point_by_point_reply_"):
            pass
        elif line.startswith("<referee_comment>"):
            referee_comment_has_started = True
            document.add_paragraph("Comments: ", style="Response Heading")
        elif line.startswith("</referee_comment>"):
            referee_comment_has_started = False
            pass
        elif line.startswith("<author_response>"):
            document.add_paragraph("Reply: ", style="Response Heading")
        elif line.startswith("</author_response>"):
            pass
        elif line.startswith("<list_of_major_changes>"):
            document.add_paragraph(
                "-----------------------------------------------------------------------",
                style="Response Heading",
            )
            document.add_paragraph("List of Major Changes", style="Response Heading")
        elif line.startswith("</list_of_major_changes>"):
            pass
        elif "</reply_letter>" in line:
            pass
        elif cover_letter_has_ended:
            if referee_comment_has_started:
                document.add_paragraph(line, style="Comment")
            else:
                document.add_paragraph(line, style="Response")
        else:
            document.add_paragraph(line, style="Cover Letter")


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--input_file", help="The input file containing the structured text")
    args = parser.parse_args()

    # Read the structured text from the file
    with open(args.input_file) as file:
        text = file.read()

    # Create a new Word document
    document = Document()

    # Apply styles to the document
    apply_styles(document)

    # Process the structured text and add content to the document
    process_text(text, document)

    # Save the document
    document.save(args.input_file.replace(".txt", ".docx"))
